# word_filler.py

# word_filler.py

from docx import Document

def fill_word_template(template_path: str, output_path: str, data: dict):
    """
    Replace placeholders like {{NAME}}, {{DOB}}, etc. in a Word document
    using the values provided in the data dictionary.
    """
    doc = Document(template_path)

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        inline = paragraph.runs
        for i in range(len(inline)):
            for key, val in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, val)

    # Replace placeholders in tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        for key, val in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in inline[i].text:
                                inline[i].text = inline[i].text.replace(placeholder, val)

    doc.save(output_path)

